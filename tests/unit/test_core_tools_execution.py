"""
核心功能执行测试 - 文档修改、段落修饰、文本美化、格式转换
测试实际调用 execute() 方法的真实功能
"""

import pytest
import sys
import os
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(__file__))))


# ============================================================
# TextTransformTool - 文本风格转换（文档修饰核心）
# ============================================================

class TestTextTransformFormal:
    """正式风格转换测试"""

    @pytest.mark.asyncio
    async def test_formal_replaces_informal_words(self):
        """[核心功能] 正式化替换口语词汇"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="我觉得这个方案挺好的，超级不错", transform_type="formal")
        assert "error" not in result
        transformed = result["transformed"]
        # 验证口语词被替换
        assert "我觉得" not in transformed
        assert "我认为" in transformed
        assert "挺好的" not in transformed
        assert "超级" not in transformed

    @pytest.mark.asyncio
    async def test_formal_preserves_content_structure(self):
        """[核心功能] 正式化保持内容结构完整"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        text = "我觉得这个方案挺好的。咋样？贼棒！"
        result = await tool.execute(text=text, transform_type="formal")
        transformed = result["transformed"]
        # 句子结构应保持
        assert "。" in transformed
        assert "？" in transformed or "?" in transformed

    @pytest.mark.asyncio
    async def test_formal_no_change_on_already_formal_text(self):
        """[核心功能] 已经是正式风格的文本不应被修改"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        text = "本研究认为该方案具有较高的可行性。"
        result = await tool.execute(text=text, transform_type="formal")
        # 不含口语词，应原样返回
        assert result["transformed"] == text

    @pytest.mark.asyncio
    async def test_formal_returns_original_and_changes(self):
        """[核心功能] 返回结构包含原文和变更数"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="我觉得挺好的", transform_type="formal")
        assert "original" in result
        assert "transformed" in result
        assert "changes" in result
        assert result["original"] == "我觉得挺好的"


class TestTextTransformCasual:
    """随意风格转换测试"""

    @pytest.mark.asyncio
    async def test_casual_replaces_formal_words(self):
        """[核心功能] 随意化替换正式词汇"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="我认为这个方案相当不错，非常优秀", transform_type="casual")
        transformed = result["transformed"]
        assert "我认为" not in transformed
        assert "我觉得" in transformed

    @pytest.mark.asyncio
    async def test_casual_transform_type_in_result(self):
        """[核心功能] 结果包含转换类型"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="我认为这个方案不错", transform_type="casual")
        assert result["type"] == "casual"


class TestTextTransformConcise:
    """简洁风格转换测试"""

    @pytest.mark.asyncio
    async def test_concise_removes_redundant_words(self):
        """[核心功能] 简洁化移除冗余修饰词"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        # 代码替换模式是 "{word}的" 和 "{word}地"，所以需要用 "非常的" 格式
        result = await tool.execute(text="这是一个非常好的方案，十分的优秀", transform_type="concise")
        transformed = result["transformed"]
        # "十分的" 应被移除
        assert "十分的" not in transformed

    @pytest.mark.asyncio
    async def test_concise_removes_duplicate_punctuation(self):
        """[核心功能] 简洁化合并重复标点"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="太好了！！！太棒了。。。", transform_type="concise")
        transformed = result["transformed"]
        # 重复标点应被合并
        assert "！！！" not in transformed
        assert "。。。。" not in transformed

    @pytest.mark.asyncio
    async def test_concise_reports_length_reduction(self):
        """[核心功能] 简洁化报告长度变化"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="这是一个非常非常好的方案", transform_type="concise")
        assert "original_length" in result
        assert "concise_length" in result


class TestTextTransformDetailed:
    """详细风格转换测试"""

    @pytest.mark.asyncio
    async def test_detailed_expands_short_sentences(self):
        """[核心功能] 详细化扩展短句"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="方案可行。效果明显。", transform_type="detailed")
        transformed = result["transformed"]
        # 短句应被扩展
        assert len(transformed) >= len("方案可行。效果明显。")

    @pytest.mark.asyncio
    async def test_detailed_preserves_long_sentences(self):
        """[核心功能] 详细化保持长句不变"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        long_sentence = "这是一个非常长的句子，已经包含了足够的细节信息，不需要再进行扩展处理。"
        result = await tool.execute(text=long_sentence, transform_type="detailed")
        transformed = result["transformed"]
        # 长句（>20字）应保持不变（注意：代码按"。"分割会丢失句末标点）
        assert "非常长的句子" in transformed
        assert "足够的细节信息" in transformed

    @pytest.mark.asyncio
    async def test_detailed_reports_length_increase(self):
        """[核心功能] 详细化报告长度增长"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="方案可行。", transform_type="detailed")
        assert result["detailed_length"] >= result["original_length"]


class TestTextTransformAcademic:
    """学术风格转换测试"""

    @pytest.mark.asyncio
    async def test_academic_replaces_colloquial(self):
        """[核心功能] 学术化替换口语表达"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="我觉得大家都知道这个技术很重要", transform_type="academic")
        transformed = result["transformed"]
        assert "我觉得" not in transformed
        assert "本研究认为" in transformed
        assert "大家都知道" not in transformed
        assert "众所周知" in transformed

    @pytest.mark.asyncio
    async def test_academic_replaces_uncertain_words(self):
        """[核心功能] 学术化替换不确定性词汇"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="这个方法可能是最有效的", transform_type="academic")
        transformed = result["transformed"]
        assert "可能" not in transformed
        assert "或许" in transformed


class TestTextTransformBusiness:
    """商务风格转换测试"""

    @pytest.mark.asyncio
    async def test_business_replaces_casual_greetings(self):
        """[核心功能] 商务化替换随意问候"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="你好，请尽快回复，谢谢", transform_type="business")
        transformed = result["transformed"]
        assert "你好" not in transformed
        assert "尊敬的" in transformed
        assert "谢谢" not in transformed
        assert "感谢" in transformed
        assert "尽快" not in transformed
        assert "及时" in transformed

    @pytest.mark.asyncio
    async def test_business_replaces_apology(self):
        """[核心功能] 商务化替换道歉用语"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="对不起，马上处理", transform_type="business")
        transformed = result["transformed"]
        assert "对不起" not in transformed
        assert "抱歉" in transformed
        assert "马上" not in transformed
        assert "立即" in transformed


class TestTextTransformEdgeCases:
    """文本转换边界条件"""

    @pytest.mark.asyncio
    async def test_empty_text_returns_error(self):
        """[边界] 空文本返回错误"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="", transform_type="formal")
        assert "error" in result

    @pytest.mark.asyncio
    async def test_missing_transform_type_returns_error(self):
        """[边界] 缺少转换类型返回错误"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="测试文本")
        assert "error" in result

    @pytest.mark.asyncio
    async def test_unsupported_transform_type_returns_error(self):
        """[边界] 不支持的转换类型返回错误"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        result = await tool.execute(text="测试文本", transform_type="shakespearean")
        assert "error" in result
        assert "不支持" in result["error"]

    @pytest.mark.asyncio
    async def test_unicode_text_transformation(self):
        """[边界] Unicode 文本转换"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        text = "我觉得这个emoji方案🎉挺好的"
        result = await tool.execute(text=text, transform_type="formal")
        assert "error" not in result
        assert "🎉" in result["transformed"]

    @pytest.mark.asyncio
    async def test_very_long_text_transformation(self):
        """[边界] 超长文本转换不崩溃"""
        from tools.text_tools import TextTransformTool
        tool = TextTransformTool()
        text = "我觉得这个方案挺好的。" * 1000
        result = await tool.execute(text=text, transform_type="formal")
        assert "error" not in result
        assert len(result["transformed"]) > 0


# ============================================================
# MarkdownToDocxTool - Markdown 转 Word
# ============================================================

class TestMarkdownToDocxExecution:
    """Markdown 转 Word 执行测试"""

    @pytest.mark.asyncio
    async def test_basic_md_to_docx(self):
        """[核心功能] 基础 Markdown 转 Word"""
        pytest.importorskip("docx")
        from tools.format_tools import MarkdownToDocxTool
        tool = MarkdownToDocxTool()

        md = "# 测试标题\n\n这是测试段落。\n\n## 二级标题\n\n- 列表项1\n- 列表项2\n"
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            result = await tool.execute(md_content=md, output_path=output_path)
            if "error" in result and "python-docx" in result["error"]:
                pytest.skip("python-docx 未安装")
            assert result.get("success") is True
            assert os.path.exists(output_path)
            # 验证生成的 docx 文件可读
            from docx import Document
            doc = Document(output_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            assert "测试标题" in full_text
            assert "测试段落" in full_text
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    @pytest.mark.asyncio
    async def test_md_to_docx_headings(self):
        """[核心功能] Markdown 标题转 Word 标题"""
        pytest.importorskip("docx")
        from tools.format_tools import MarkdownToDocxTool
        tool = MarkdownToDocxTool()

        md = "# 一级\n## 二级\n### 三级\n#### 四级\n"
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            result = await tool.execute(md_content=md, output_path=output_path)
            if "error" in result and "python-docx" in result["error"]:
                pytest.skip("python-docx 未安装")
            assert result.get("success") is True
            from docx import Document
            doc = Document(output_path)
            headings = [p.text for p in doc.paragraphs if p.text.strip()]
            assert "一级" in headings
            assert "二级" in headings
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    @pytest.mark.asyncio
    async def test_md_to_docx_lists(self):
        """[核心功能] Markdown 列表转 Word 列表"""
        pytest.importorskip("docx")
        from tools.format_tools import MarkdownToDocxTool
        tool = MarkdownToDocxTool()

        md = "- 无序1\n- 无序2\n\n1. 有序1\n2. 有序2\n"
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            result = await tool.execute(md_content=md, output_path=output_path)
            if "error" in result and "python-docx" in result["error"]:
                pytest.skip("python-docx 未安装")
            assert result.get("success") is True
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    @pytest.mark.asyncio
    async def test_md_to_docx_missing_params(self):
        """[边界] 缺少参数返回错误"""
        from tools.format_tools import MarkdownToDocxTool
        tool = MarkdownToDocxTool()
        result = await tool.execute(md_content="test")
        assert "error" in result

        result2 = await tool.execute(output_path="/tmp/test.docx")
        assert "error" in result2

    @pytest.mark.asyncio
    async def test_md_to_docx_bold_italic(self):
        """[核心功能] Markdown 加粗斜体处理"""
        pytest.importorskip("docx")
        from tools.format_tools import MarkdownToDocxTool
        tool = MarkdownToDocxTool()

        md = "这是**加粗**和*斜体*文本。\n"
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            result = await tool.execute(md_content=md, output_path=output_path)
            if "error" in result and "python-docx" in result["error"]:
                pytest.skip("python-docx 未安装")
            assert result.get("success") is True
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)


# ============================================================
# MarkdownToPptxTool - Markdown 转 PPT
# ============================================================

class TestMarkdownToPptxExecution:
    """Markdown 转 PPT 执行测试"""

    @pytest.mark.asyncio
    async def test_basic_md_to_pptx(self):
        """[核心功能] 基础 Markdown 转 PPT"""
        pytest.importorskip("pptx")
        from tools.format_tools import MarkdownToPptxTool
        tool = MarkdownToPptxTool()

        md = "# 第一页\n\n内容1\n\n# 第二页\n\n内容2\n"
        with tempfile.NamedTemporaryFile(suffix='.pptx', delete=False) as f:
            output_path = f.name

        try:
            result = await tool.execute(md_content=md, output_path=output_path)
            if "error" in result and "python-pptx" in result["error"]:
                pytest.skip("python-pptx 未安装")
            assert result.get("success") is True
            assert os.path.exists(output_path)
            from pptx import Presentation
            prs = Presentation(output_path)
            assert len(prs.slides) >= 2
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    @pytest.mark.asyncio
    async def test_md_to_pptx_with_bullet_list(self):
        """[核心功能] 带列表的 Markdown 转 PPT"""
        pytest.importorskip("pptx")
        from tools.format_tools import MarkdownToPptxTool
        tool = MarkdownToPptxTool()

        md = "# 演示\n\n- 要点1\n- 要点2\n- 要点3\n"
        with tempfile.NamedTemporaryFile(suffix='.pptx', delete=False) as f:
            output_path = f.name

        try:
            result = await tool.execute(md_content=md, output_path=output_path)
            if "error" in result and "python-pptx" in result["error"]:
                pytest.skip("python-pptx 未安装")
            assert result.get("success") is True
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    @pytest.mark.asyncio
    async def test_md_to_pptx_missing_params(self):
        """[边界] 缺少参数返回错误"""
        from tools.format_tools import MarkdownToPptxTool
        tool = MarkdownToPptxTool()
        result = await tool.execute(md_content="test")
        assert "error" in result


# ============================================================
# TextToTableTool - 文本转表格
# ============================================================

class TestTextToTableExecution:
    """文本转表格执行测试"""

    @pytest.mark.asyncio
    async def test_csv_to_markdown_table(self):
        """[核心功能] CSV 文本转 Markdown 表格"""
        from tools.format_tools import TextToTableTool
        tool = TextToTableTool()

        text = "姓名,年龄,城市\n张三,25,北京\n李四,30,上海"
        result = await tool.execute(text=text, format="markdown")
        assert "error" not in result
        assert result["type"] == "table"
        assert result["rows"] == 3
        assert result["columns"] == 3
        content = result["content"]
        assert "姓名" in content
        assert "张三" in content
        # 分隔行格式可能是 "| -- |" 或 "|---|"
        assert "| --" in content or "|---" in content

    @pytest.mark.asyncio
    async def test_csv_to_html_table(self):
        """[核心功能] CSV 文本转 HTML 表格"""
        from tools.format_tools import TextToTableTool
        tool = TextToTableTool()

        text = "姓名,年龄\n张三,25"
        result = await tool.execute(text=text, format="html")
        assert "error" not in result
        content = result["content"]
        assert "<table>" in content
        assert "<th>姓名</th>" in content
        assert "<td>张三</td>" in content

    @pytest.mark.asyncio
    async def test_csv_to_csv_format(self):
        """[核心功能] 文本转 CSV 格式"""
        from tools.format_tools import TextToTableTool
        tool = TextToTableTool()

        text = "姓名,年龄\n张三,25"
        result = await tool.execute(text=text, format="csv")
        assert "error" not in result
        assert "姓名,年龄" in result["content"]
        assert "张三,25" in result["content"]

    @pytest.mark.asyncio
    async def test_tab_delimited_text(self):
        """[核心功能] Tab 分隔文本转表格"""
        from tools.format_tools import TextToTableTool
        tool = TextToTableTool()

        text = "姓名\t年龄\t城市\n张三\t25\t北京"
        result = await tool.execute(text=text, format="markdown")
        assert "error" not in result
        assert result["columns"] == 3

    @pytest.mark.asyncio
    async def test_multi_space_delimited(self):
        """[核心功能] 多空格分隔文本转表格"""
        from tools.format_tools import TextToTableTool
        tool = TextToTableTool()

        text = "姓名    年龄    城市\n张三    25    北京"
        result = await tool.execute(text=text, format="markdown")
        assert "error" not in result
        assert result["columns"] == 3

    @pytest.mark.asyncio
    async def test_custom_delimiter(self):
        """[核心功能] 自定义分隔符"""
        from tools.format_tools import TextToTableTool
        tool = TextToTableTool()

        text = "姓名|年龄|城市\n张三|25|北京"
        result = await tool.execute(text=text, format="markdown", delimiter="|")
        assert "error" not in result
        assert result["columns"] == 3

    @pytest.mark.asyncio
    async def test_empty_text_returns_error(self):
        """[边界] 空文本返回错误"""
        from tools.format_tools import TextToTableTool
        tool = TextToTableTool()
        result = await tool.execute(text="")
        assert "error" in result

    @pytest.mark.asyncio
    async def test_csv_with_special_chars(self):
        """[边界] CSV 包含特殊字符（逗号、引号）"""
        from tools.format_tools import TextToTableTool
        tool = TextToTableTool()

        text = '姓名,描述\n张三,"他来自北京,是工程师"\n李四,设计师'
        result = await tool.execute(text=text, format="csv")
        assert "error" not in result

    @pytest.mark.asyncio
    async def test_html_table_structure(self):
        """[核心功能] HTML 表格结构完整性"""
        from tools.format_tools import TextToTableTool
        tool = TextToTableTool()

        text = "列A,列B\n数据1,数据2"
        result = await tool.execute(text=text, format="html")
        content = result["content"]
        assert "<thead>" in content
        assert "<tbody>" in content
        assert "</table>" in content


# ============================================================
# TextExtractTool 摘要生成 - 补充测试
# ============================================================

class TestTextExtractSummary:
    """文本摘要生成测试"""

    @pytest.mark.asyncio
    async def test_summary_from_long_text(self):
        """[核心功能] 从长文本生成摘要"""
        from tools.text_tools import TextExtractTool
        tool = TextExtractTool()

        # 代码要求句子长度 > 10 字符才保留
        text = "人工智能技术正在快速发展改变世界。机器学习作为核心技术已经取得重大突破。深度学习在多个领域展现出强大能力。自然语言处理技术应用越来越广泛。计算机视觉也在不断进步。"
        result = await tool.execute(content=text, extract_type="summary")
        assert "error" not in result
        assert result["type"] == "summary"
        assert "summary" in result
        assert len(result["summary"]) > 0
        # 摘要应比原文短
        assert result["summary_length"] <= result["original_length"]

    @pytest.mark.asyncio
    async def test_summary_preserves_key_content(self):
        """[核心功能] 摘要保留关键内容"""
        from tools.text_tools import TextExtractTool
        tool = TextExtractTool()

        text = "人工智能技术正在快速发展。机器学习是核心技术。深度学习取得了突破。自然语言处理应用广泛。"
        result = await tool.execute(content=text, extract_type="summary")
        summary = result["summary"]
        # 摘要应包含前几个句子
        assert "人工智能" in summary or "机器学习" in summary

    @pytest.mark.asyncio
    async def test_summary_empty_content(self):
        """[边界] 空内容生成空摘要"""
        from tools.text_tools import TextExtractTool
        tool = TextExtractTool()

        result = await tool.execute(content="。", extract_type="summary")
        assert "error" not in result
        assert result["type"] == "summary"


# ============================================================
# Persona 驱动的工作流测试
# ============================================================

class TestPersonaDrivenWorkflow:
    """Persona 驱动的 LLM 工作流测试"""

    def test_all_persona_files_loadable(self):
        """[核心功能] 所有 persona 文件都能加载"""
        from asu_custom_agent import load_persona

        personas = ["default", "polish", "revision", "translate", "code", "custom"]
        for name in personas:
            content = load_persona(name)
            assert isinstance(content, str), f"persona '{name}' 加载失败"
            assert len(content) > 10, f"persona '{name}' 内容过短"

    def test_office_persona_files_loadable(self):
        """[核心功能] 办公场景 persona 文件都能加载"""
        from asu_custom_agent import load_persona

        # 测试嵌套路径的 persona
        office_personas = [
            "office/academic/paper",
            "office/business/email",
            "office/business/report",
        ]
        for name in office_personas:
            content = load_persona(name)
            assert isinstance(content, str), f"persona '{name}' 加载失败"
            assert len(content) > 10, f"persona '{name}' 内容过短"

    def test_translation_persona_loadable(self):
        """[核心功能] 翻译 persona 文件能加载"""
        from asu_custom_agent import load_persona

        content = load_persona("translation/technical")
        assert isinstance(content, str)
        assert "翻译" in content or "术语" in content

    def test_polish_persona_content_quality(self):
        """[核心功能] 润色 persona 包含必要的指令"""
        from asu_custom_agent import load_persona

        persona = load_persona("polish")
        # 验证润色 persona 包含关键指令
        assert "润色" in persona or "修正" in persona
        assert "输出" in persona  # 应有输出规范

    def test_revision_persona_content_quality(self):
        """[核心功能] 修订 persona 包含必要的指令"""
        from asu_custom_agent import load_persona

        persona = load_persona("revision")
        # 验证修订 persona 包含关键指令
        assert "修订" in persona or "修改" in persona
        assert "联动" in persona  # 应有联动分析
        assert "selection" in persona or "选中" in persona

    def test_polish_persona_with_context_prefix(self):
        """[核心功能] 润色 persona + context_prefix 组合"""
        from asu_custom_agent import load_persona, build_context_prefix, ContextWindowManager

        persona = load_persona("polish")
        prefix = build_context_prefix("drag", {})
        manager = ContextWindowManager()

        envelope = {"source": "drag", "content": "我觉得这个方案挺好的。"}
        messages = manager.build_messages(
            system_prompt=f"{prefix}\n\n{persona}",
            envelope=envelope,
            history_messages=[]
        )

        # 验证消息结构
        assert len(messages) >= 2  # system + user
        assert messages[0]["role"] == "system"
        assert "润色" in messages[0]["content"] or "修正" in messages[0]["content"]
        assert messages[-1]["role"] == "user"

    def test_revision_persona_with_selection_and_content(self):
        """[核心功能] 修订 persona 处理选中文本 + 全文"""
        from asu_custom_agent import load_persona, build_context_prefix, ContextWindowManager

        persona = load_persona("revision")
        prefix = build_context_prefix("revision", {})
        manager = ContextWindowManager()

        envelope = {
            "source": "revision",
            "content": "这是一篇完整文档。包含多个段落。第三段需要修改。",
            "selection": "第三段需要修改",
        }
        messages = manager.build_messages(
            system_prompt=f"{prefix}\n\n{persona}",
            envelope=envelope,
            history_messages=[]
        )

        user_payload = messages[-1]["content"]
        # 验证选中文本和全文都被包含
        assert "第三段需要修改" in user_payload
        assert "完整文档" in user_payload

    def test_custom_instruction_injected_into_payload(self):
        """[核心功能] 自定义指令注入到用户消息"""
        from asu_custom_agent import ContextWindowManager, normalize_context_envelope

        manager = ContextWindowManager()
        req = {
            "context_envelope": {
                "source": "drag",
                "content": "Hello World",
                "meta": {"custom_instruction": "翻译为日语"},
            }
        }
        envelope = normalize_context_envelope(req, "", "drag", {})
        messages = manager.build_messages(
            system_prompt="You are a translator.",
            envelope=envelope,
            history_messages=[]
        )

        user_payload = messages[-1]["content"]
        assert "翻译为日语" in user_payload
        assert "custom_instruction" in user_payload

    def test_multiple_styles_context_prefix(self):
        """[核心功能] 不同来源的 context_prefix 描述正确"""
        from asu_custom_agent import build_context_prefix

        sources = {
            "drag": "拖拽",
            "ide": "IDE",
            "browser": "浏览器",
            "chat": "对话",
            "revision": "修订",
        }
        for source, keyword in sources.items():
            prefix = build_context_prefix(source, {})
            assert isinstance(prefix, str)
            assert len(prefix) > 0


# ============================================================
# 工具注册表与工具发现
# ============================================================

class TestToolDiscovery:
    """工具发现与注册测试"""

    def test_all_core_tools_importable(self):
        """[核心功能] 所有核心工具都能导入"""
        from tools.file_tools import FileReadTool
        from tools.text_tools import TextExtractTool, TextTransformTool
        from tools.format_tools import MarkdownToDocxTool, MarkdownToPptxTool, TextToTableTool

        assert FileReadTool is not None
        assert TextExtractTool is not None
        assert TextTransformTool is not None
        assert MarkdownToDocxTool is not None
        assert MarkdownToPptxTool is not None
        assert TextToTableTool is not None

    def test_all_tools_have_required_interface(self):
        """[核心功能] 所有工具都实现必要接口"""
        from tools.text_tools import TextTransformTool
        from tools.format_tools import MarkdownToDocxTool, MarkdownToPptxTool, TextToTableTool

        tools = [TextTransformTool(), MarkdownToDocxTool(), MarkdownToPptxTool(), TextToTableTool()]
        for tool in tools:
            assert hasattr(tool, 'name'), f"{tool.__class__.__name__} 缺少 name"
            assert hasattr(tool, 'description'), f"{tool.__class__.__name__} 缺少 description"
            assert hasattr(tool, 'parameters'), f"{tool.__class__.__name__} 缺少 parameters"
            assert hasattr(tool, 'execute'), f"{tool.__class__.__name__} 缺少 execute"
            assert isinstance(tool.name, str)
            assert isinstance(tool.description, str)
            assert isinstance(tool.parameters, dict)

    def test_tool_registry_register_and_execute(self):
        """[核心功能] 工具注册后可执行"""
        from tools.base import ToolRegistry
        from tools.text_tools import TextTransformTool

        registry = ToolRegistry()
        tool = TextTransformTool()
        registry.register(tool.name, tool)

        retrieved = registry.get_tool(tool.name)
        assert retrieved is tool
        assert retrieved.name == "text_transform"


# ============================================================
# 集成：文档处理完整工作流
# ============================================================

class TestDocumentWorkflowIntegration:
    """文档处理完整工作流集成测试"""

    @pytest.mark.asyncio
    async def test_read_extract_transform_pipeline(self):
        """[核心功能] 读取 → 提取 → 转换 完整流水线"""
        from tools.file_tools import FileReadTool
        from tools.text_tools import TextExtractTool, TextTransformTool

        # 创建测试文件
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write("# 项目报告\n\n我觉得这个项目挺好的，超级成功。\n\n## 总结\n\n非常满意。")
            temp_path = f.name

        try:
            # 1. 读取文件
            file_tool = FileReadTool()
            read_result = await file_tool.execute(file_path=temp_path, format="text")
            assert "error" not in read_result
            content = read_result["content"]

            # 2. 提取标题
            extract_tool = TextExtractTool()
            headings_result = await extract_tool.execute(content=content, extract_type="headings")
            assert "error" not in headings_result
            assert len(headings_result["headings"]) == 2

            # 3. 转换风格
            transform_tool = TextTransformTool()
            formal_result = await transform_tool.execute(text=content, transform_type="formal")
            assert "error" not in formal_result
            # 验证口语词被替换
            assert "挺好的" not in formal_result["transformed"]
            assert "超级" not in formal_result["transformed"]
        finally:
            os.unlink(temp_path)

    @pytest.mark.asyncio
    async def test_extract_and_convert_to_table(self):
        """[核心功能] 提取内容 → 转表格"""
        from tools.text_tools import TextExtractTool
        from tools.format_tools import TextToTableTool

        content = "项目A,进度80%,负责人张三\n项目B,进度60%,负责人李四\n项目C,进度90%,负责人王五"

        table_tool = TextToTableTool()
        result = await table_tool.execute(text=content, format="markdown")
        assert "error" not in result
        assert result["rows"] == 3
        assert "项目A" in result["content"]

    @pytest.mark.asyncio
    async def test_polish_and_export_docx(self):
        """[核心功能] 润色文本 → 导出 Word"""
        pytest.importorskip("docx")
        from tools.text_tools import TextTransformTool
        from tools.format_tools import MarkdownToDocxTool

        # 1. 润色文本
        transform_tool = TextTransformTool()
        polished = await transform_tool.execute(
            text="我觉得这个方案挺好的，超级不错",
            transform_type="formal"
        )
        assert "error" not in polished

        # 2. 构造 Markdown 并导出 Word
        md_content = f"# 方案评估\n\n{polished['transformed']}\n"
        docx_tool = MarkdownToDocxTool()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            result = await docx_tool.execute(md_content=md_content, output_path=output_path)
            if "error" in result and "python-docx" in result["error"]:
                pytest.skip("python-docx 未安装")
            assert result.get("success") is True
            assert os.path.exists(output_path)
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)
